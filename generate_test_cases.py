"""
Generates Maskan Portal Test Cases Word Document
Run: python generate_test_cases.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Colour palette ──────────────────────────────────────────────────────────
C_DARK       = RGBColor(0x1E, 0x29, 0x3B)   # slate-800  – headings
C_PRIMARY    = RGBColor(0x22, 0x63, 0xEB)   # blue-600   – section banners
C_ACCENT     = RGBColor(0x0E, 0xA5, 0xE9)   # sky-500    – sub-headers
C_SUCCESS    = RGBColor(0x16, 0xA3, 0x4A)   # green-600  – pass column
C_WARN       = RGBColor(0xD9, 0x77, 0x06)   # amber-600  – notes
C_LIGHT      = RGBColor(0xF1, 0xF5, 0xF9)   # slate-100  – table header bg
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_BORDER     = RGBColor(0xE2, 0xE8, 0xF0)   # slate-200

# ─── Helpers ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "E2E8F0")
        tcBorders.append(border)
    tcPr.append(tcBorders)

def cell_para(cell, text, bold=False, size=9, color=None, align=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p   = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or C_DARK
    return p

def add_step_run(cell, steps: list[str]):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p = cell.paragraphs[0]
    for idx, step in enumerate(steps, 1):
        if idx > 1:
            p.add_run("\n")
        r = p.add_run(f"{idx}. {step}")
        r.font.size = Pt(9)
        r.font.color.rgb = C_DARK

def add_result_run(cell, expected: str):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p   = cell.paragraphs[0]
    run = p.add_run(expected)
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x14, 0x53, 0x2D)   # green-900

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)

# ─── Document setup ───────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.page_width  = Inches(11.69)   # A4 landscape
section.page_height = Inches(8.27)
section.left_margin   = Inches(0.6)
section.right_margin  = Inches(0.6)
section.top_margin    = Inches(0.6)
section.bottom_margin = Inches(0.6)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ─── Cover page ───────────────────────────────────────────────────────────────
def cover_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("MASKAN RENTAL PORTAL")
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = C_PRIMARY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Test Case Specification")
    r2.font.size = Pt(18); r2.font.color.rgb = C_DARK

    doc.add_paragraph()
    info = [
        ("Document Type:",  "User Acceptance Test Cases"),
        ("Portal URL:",      "http://localhost:8080"),
        ("Backend URL:",     "http://localhost:8000"),
        ("Version:",         "1.0"),
        ("Prepared by:",     "Maskan Development Team"),
        ("Date:",            datetime.date.today().strftime("%d %B %Y")),
    ]
    t = doc.add_table(rows=len(info), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info):
        t.rows[i].cells[0].width = Cm(5)
        t.rows[i].cells[1].width = Cm(10)
        cell_para(t.rows[i].cells[0], k, bold=True, size=11)
        cell_para(t.rows[i].cells[1], v, size=11)

    doc.add_page_break()

# ─── Section banner ───────────────────────────────────────────────────────────
def section_banner(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(f"  SECTION {num:02d} — {title.upper()}  ")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = C_WHITE
    rPr = r._r.get_or_add_rPr()
    # background shading via paragraph shade
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "2263EB")
    pPr.append(shd)

# ─── Table builder ────────────────────────────────────────────────────────────
COLS    = ["TC ID", "Feature", "Test Case Title", "Pre-conditions", "Test Steps", "Expected Result", "Pass/Fail", "Notes"]
WIDTHS  = [1.6,      2.2,       3.8,               3.0,              6.5,          5.5,                1.4,          2.5]

def make_table(doc, cases):
    table = doc.add_table(rows=1, cols=len(COLS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, (col, w) in enumerate(zip(COLS, WIDTHS)):
        cell = hdr.cells[i]
        cell.width = Cm(w)
        set_cell_bg(cell, "1E293B")
        cell_para(cell, col, bold=True, size=9, color=C_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for idx, tc in enumerate(cases):
        row  = table.add_row()
        bg   = "FFFFFF" if idx % 2 == 0 else "F8FAFC"

        for i, cell in enumerate(row.cells):
            set_cell_bg(cell, bg)
            set_cell_border(cell)

        cell_para(row.cells[0], tc["id"],    bold=True, size=9, color=C_PRIMARY)
        cell_para(row.cells[1], tc["feature"],           size=9)
        cell_para(row.cells[2], tc["title"],  bold=True, size=9)
        cell_para(row.cells[3], tc["pre"],               size=9)
        add_step_run(row.cells[4], tc["steps"])
        add_result_run(row.cells[5], tc["expected"])
        cell_para(row.cells[6], "☐ Pass\n☐ Fail", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(row.cells[7], tc.get("notes", ""), size=8, color=C_WARN)

    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════════════════════════
# TEST CASES DATA
# ═══════════════════════════════════════════════════════════════════════════════

sections = []

# ── S01: Home Page ─────────────────────────────────────────────────────────────
sections.append(("Home Page / Landing", [
    {
        "id": "TC-001", "feature": "Home – Navigation",
        "title": "Top navigation links load correct pages",
        "pre": "Portal running at localhost:8080",
        "steps": [
            "Open http://localhost:8080",
            'Click "Search" in the top nav',
            "Click browser back",
            'Click "Explore Areas" in the top nav',
            "Click browser back",
            'Click "AI Advisor" in the top nav',
            "Click browser back",
        ],
        "expected": "Each nav link navigates to /search, /areas, and /advisor respectively. Back button returns to home each time.",
        "notes": "Verify logo click also returns to home",
    },
    {
        "id": "TC-002", "feature": "Home – City cards",
        "title": "City cards display listing counts from live data",
        "pre": "At least one published property per city",
        "steps": [
            "Open http://localhost:8080",
            "Scroll to the city selector row",
            "Note the listing counts shown under each city (Riyadh, Jeddah, Dammam, Khobar, Madinah)",
            'Click the "Riyadh" city card',
        ],
        "expected": "City cards show non-zero listing counts loaded from the backend. Clicking a city navigates to /search filtered for that city.",
        "notes": "Counts must match backend DB, not be hardcoded",
    },
    {
        "id": "TC-003", "feature": "Home – Featured areas",
        "title": "Featured areas section shows AI scores",
        "pre": "Area Intelligence data seeded in DB",
        "steps": [
            "Open http://localhost:8080",
            "Scroll to the Featured Areas section",
            "Verify each area card shows: Area name, Property Score (0–100), Family Score, Rental Value",
            "Click on any area card",
        ],
        "expected": "Area cards display real scores from AreaIntelligence table. Clicking an area card opens that area's detail on /areas.",
        "notes": "",
    },
    {
        "id": "TC-004", "feature": "Home – Featured properties",
        "title": "Featured properties carousel shows live listings",
        "pre": "Published properties exist in DB",
        "steps": [
            "Open http://localhost:8080",
            "Scroll to Featured Properties section",
            "Verify at least 3 property cards are visible",
            "Click 'View details' on any property card",
        ],
        "expected": "Property cards loaded from API. Each shows title, district, price, bedrooms. View details opens /property/$id.",
        "notes": "",
    },
    {
        "id": "TC-005", "feature": "Home – Quick AI questions",
        "title": "Quick question chips navigate to AI Advisor",
        "pre": "Home page loaded",
        "steps": [
            "Scroll to the AI Quick Questions section",
            "Click any suggested question chip",
        ],
        "expected": "Page navigates to /advisor and the question is pre-filled or sent automatically.",
        "notes": "",
    },
]))

# ── S02: Authentication ─────────────────────────────────────────────────────────
sections.append(("Authentication – Sign Up & Sign In", [
    {
        "id": "TC-006", "feature": "Auth – Sign Up",
        "title": "New user can register with email and password",
        "pre": "Email not already registered",
        "steps": [
            "Open http://localhost:8080/auth",
            'Click "Sign up" tab',
            "Enter Full Name: Test User",
            "Enter Email: testuser@maskan.test",
            "Enter Password: Test@1234",
            "Check Terms & Conditions checkbox",
            'Click "Create account" button',
        ],
        "expected": "Account created. User is redirected to home or dashboard. Auth token stored in session. User name visible in nav.",
        "notes": "Check: error shown if email already exists",
    },
    {
        "id": "TC-007", "feature": "Auth – Sign Up Validation",
        "title": "Sign up form validates required fields",
        "pre": "On /auth page, Sign up tab active",
        "steps": [
            "Leave all fields empty, click Create account",
            "Enter email only (no password), click Create account",
            "Enter a weak password (less than 6 chars), click Create account",
            "Enter valid email + password, leave T&C unchecked, click Create account",
        ],
        "expected": "Each step shows an inline error message. Form does not submit until all required fields are filled and T&C is checked.",
        "notes": "",
    },
    {
        "id": "TC-008", "feature": "Auth – Sign In",
        "title": "Existing user can sign in with correct credentials",
        "pre": "Test user account already created (TC-006)",
        "steps": [
            "Open http://localhost:8080/auth",
            'Ensure "Sign in" tab is active',
            "Enter Email: testuser@maskan.test",
            "Enter Password: Test@1234",
            'Click "Sign in" button',
        ],
        "expected": "User is authenticated. Redirected to home or previous page. Nav shows logged-in state.",
        "notes": "",
    },
    {
        "id": "TC-009", "feature": "Auth – Sign In Validation",
        "title": "Sign in shows error for wrong credentials",
        "pre": "On /auth page",
        "steps": [
            "Enter Email: testuser@maskan.test",
            "Enter Password: WrongPass123",
            'Click "Sign in"',
        ],
        "expected": "Error message displayed: incorrect email or password. Form remains on the page.",
        "notes": "",
    },
    {
        "id": "TC-010", "feature": "Auth – Sign Out",
        "title": "Signed-in user can sign out",
        "pre": "User is signed in",
        "steps": [
            "Locate Sign Out option in navigation or user menu",
            "Click Sign Out",
        ],
        "expected": "User is signed out. Auth token cleared. Nav returns to guest state. Protected pages (Saved, Mediator) redirect to /auth.",
        "notes": "",
    },
    {
        "id": "TC-011", "feature": "Auth – Protected routes",
        "title": "Unauthenticated access to saved properties is blocked",
        "pre": "User is NOT signed in",
        "steps": [
            "Navigate directly to http://localhost:8080/saved",
        ],
        "expected": "User is redirected to /auth with a prompt to sign in.",
        "notes": "Also verify /mediator requires auth",
    },
]))

# ── S03: Search ────────────────────────────────────────────────────────────────
sections.append(("Property Search & Filtering", [
    {
        "id": "TC-012", "feature": "Search – Page load",
        "title": "Search page loads with all published properties",
        "pre": "Published properties exist in DB",
        "steps": [
            "Navigate to http://localhost:8080/search",
            "Wait for properties to load",
            "Count number of property cards shown",
        ],
        "expected": "Page loads without errors. All published properties are shown with images, title, district, price. Loading spinner shown briefly then replaced by results.",
        "notes": "Check browser console for errors",
    },
    {
        "id": "TC-013", "feature": "Search – City filter",
        "title": "City filter narrows results to selected city",
        "pre": "Search page loaded with results",
        "steps": [
            "In the filter panel, open the City dropdown",
            "Select 'Riyadh'",
            'Click "Apply filters"',
            "Count properties shown",
            "Verify each card shows 'Riyadh' in the location",
        ],
        "expected": "Only properties in Riyadh are displayed. Count matches number of Riyadh properties in DB.",
        "notes": "",
    },
    {
        "id": "TC-014", "feature": "Search – District filter",
        "title": "District dropdown populates based on selected city",
        "pre": "City filter set to 'Riyadh'",
        "steps": [
            "Open the District dropdown",
            "Observe the options available",
            "Select 'Al Narjis'",
            'Click "Apply filters"',
        ],
        "expected": "District list shows only districts in Riyadh. Selecting Al Narjis shows only Al Narjis properties.",
        "notes": "",
    },
    {
        "id": "TC-015", "feature": "Search – Price range filter",
        "title": "Min/max rent filters correctly bound results",
        "pre": "Search page loaded",
        "steps": [
            "Enter Min Rent: 50000 and Max Rent: 100000 (SAR/year)",
            'Click "Apply filters"',
            "Check each result card's annual rent",
        ],
        "expected": "All displayed properties have annual rent between SAR 50,000 and SAR 100,000.",
        "notes": "SAR amounts are annual",
    },
    {
        "id": "TC-016", "feature": "Search – Bedroom filter",
        "title": "Bedroom segmented control filters by minimum bedrooms",
        "pre": "Search page loaded",
        "steps": [
            "Click '3+' in the Bedrooms control",
            'Click "Apply filters"',
            "Check bedroom count on each result card",
        ],
        "expected": "All displayed properties have 3 or more bedrooms.",
        "notes": "",
    },
    {
        "id": "TC-017", "feature": "Search – Property type filter",
        "title": "Property type dropdown filters correctly",
        "pre": "Search page loaded",
        "steps": [
            "Open Property Type dropdown",
            "Select 'Villa'",
            'Click "Apply filters"',
        ],
        "expected": "Only Villa type properties shown in results.",
        "notes": "",
    },
    {
        "id": "TC-018", "feature": "Search – Amenities filter",
        "title": "Amenity toggles (Gym, Pool, Parking, Balcony) filter results",
        "pre": "Search page loaded",
        "steps": [
            "Toggle ON the 'Gym' amenity filter",
            'Click "Apply filters"',
            "Open one result card and view its property detail page",
            "Verify Gym is listed as an amenity",
        ],
        "expected": "Only properties with Gym amenity shown. Property detail confirms gym availability.",
        "notes": "",
    },
    {
        "id": "TC-019", "feature": "Search – Reset filters",
        "title": "Reset button clears all filters and shows all properties",
        "pre": "Multiple filters are applied",
        "steps": [
            "Apply city = Riyadh, bedrooms = 3+, rent max = 80000",
            "Note the filtered result count",
            'Click "Reset"',
        ],
        "expected": "All filter controls return to default (Any). Result count returns to full property list count.",
        "notes": "",
    },
    {
        "id": "TC-020", "feature": "Search – Sorting",
        "title": "Sort dropdown reorders results correctly",
        "pre": "Search page loaded with multiple results",
        "steps": [
            "Open Sort dropdown",
            "Select 'Price: low to high'",
            "Note prices of first 3 cards",
            "Select 'Price: high to low'",
            "Note prices of first 3 cards again",
        ],
        "expected": "Results reorder. Low-to-high shows cheapest first. High-to-low shows most expensive first.",
        "notes": "",
    },
    {
        "id": "TC-021", "feature": "Search – Property Score",
        "title": "Property Score shows on cards and is within valid range",
        "pre": "Search page loaded with results",
        "steps": [
            "View property cards on search results",
            "Check the 'Score' badge on at least 5 cards",
        ],
        "expected": "Every card shows a Property Score between 55 and 97. Score is NOT hardcoded (different cards show different scores).",
        "notes": "Was previously showing fake formula",
    },
    {
        "id": "TC-022", "feature": "Search – Save property",
        "title": "Logged-in user can save a property from search results",
        "pre": "User is signed in",
        "steps": [
            "On search results, click the heart (save) icon on any property card",
            "Observe heart icon change",
            "Navigate to /saved",
        ],
        "expected": "Heart icon fills/changes color to indicate saved. Property appears in /saved page.",
        "notes": "If not logged in, should redirect to /auth",
    },
    {
        "id": "TC-023", "feature": "Search – Grid/List view",
        "title": "View toggle switches between grid and list layouts",
        "pre": "Search page loaded with results",
        "steps": [
            "Click the List view icon in the results toolbar",
            "Observe layout change",
            "Click the Grid view icon",
        ],
        "expected": "List view shows properties in a single column with more text detail. Grid view shows cards in multi-column layout.",
        "notes": "",
    },
    {
        "id": "TC-024", "feature": "Search – Compare selection",
        "title": "Compare button adds property to compare bar",
        "pre": "Search page loaded",
        "steps": [
            "Click 'Compare' on the first property card",
            "Click 'Compare' on a second property card",
            "Observe the sticky compare bar at the bottom",
            "Click 'Compare now' in the bar",
        ],
        "expected": "Compare bar appears and shows '2 selected'. Clicking Compare now navigates to /compare with both properties pre-loaded.",
        "notes": "Max 3 properties can be selected",
    },
]))

# ── S04: Property Detail ───────────────────────────────────────────────────────
sections.append(("Property Detail Page", [
    {
        "id": "TC-025", "feature": "Property – Page load",
        "title": "Property detail page loads all sections correctly",
        "pre": "Published property exists in DB. User on search results.",
        "steps": [
            "Click 'View details' on any property from search",
            "Wait for page to fully load",
            "Scroll through the entire page from top to bottom",
        ],
        "expected": "Page shows: gallery, title, rent, property facts (beds/baths/area), Rental Intelligence, Fair Rent Analysis, Area Summary scores, AI Summary, and Comparable Listings. No blank sections.",
        "notes": "Check browser console for 404/500 errors",
    },
    {
        "id": "TC-026", "feature": "Property – Gallery",
        "title": "Image gallery navigation works correctly",
        "pre": "Property detail page loaded with images",
        "steps": [
            "Click next/prev arrows or thumbnail images",
            "Verify image counter updates (e.g. '2 / 5')",
            "Click last thumbnail (floor plan)",
        ],
        "expected": "Gallery cycles through all images. Counter shows current position. Floor plan renders as an overlay on the last image.",
        "notes": "",
    },
    {
        "id": "TC-027", "feature": "Property – Rental Intelligence",
        "title": "Rental Intelligence section shows composite score and breakdown",
        "pre": "Property detail page loaded",
        "steps": [
            "Scroll to 'Rental Intelligence' section",
            "Note the composite score (0–100)",
            "Check the breakdown bars: Price Fairness, Area Quality, Amenities, Commute, Family Suitability",
        ],
        "expected": "Score ring shows a number 55–97. Breakdown bars are filled proportionally. Verdict text shown (Excellent / Strong / Fair / Below average).",
        "notes": "",
    },
    {
        "id": "TC-028", "feature": "Property – Fair Rent Analysis",
        "title": "Fair rent analysis shows market comparison",
        "pre": "Property detail page loaded. Area intelligence data exists for this district.",
        "steps": [
            "Scroll to 'Fair Rent Analysis' section",
            "Note the current monthly rent shown",
            "Note the fair market range (SAR min–max/month)",
            "Check the verdict badge (Fair Price / Slightly Overpriced / Below Market)",
        ],
        "expected": "Section shows the property's monthly rent versus computed fair range. Rent slider visually positions current rent within the range.",
        "notes": "",
    },
    {
        "id": "TC-029", "feature": "Property – Area scores",
        "title": "Area summary shows 5 score cards from live data",
        "pre": "Area Intelligence seeded for this property's district",
        "steps": [
            "Scroll to 'Area Summary' section",
            "Check 5 score cards: Area Score, School Score, Traffic Score, Healthcare Score, Family Score",
        ],
        "expected": "All 5 scores populated from AreaIntelligence DB record. No score is 0 or missing (unless no data for district).",
        "notes": "",
    },
    {
        "id": "TC-030", "feature": "Property – Nearby places",
        "title": "Nearby places show schools, hospitals, mosques, malls, parks",
        "pre": "Area intelligence with place data exists for this district",
        "steps": [
            "Scroll to Nearby Places section",
            "Expand or view each category tab",
        ],
        "expected": "Each category (Schools, Hospitals, Mosques, Shopping Malls, Parks) shows up to 3 places with name, distance in km, and rating where available.",
        "notes": "",
    },
    {
        "id": "TC-031", "feature": "Property – Contact landlord",
        "title": "Contact landlord modal opens and accepts input",
        "pre": "Property detail page loaded",
        "steps": [
            "Click 'Contact landlord' button in the right sidebar",
            "Modal opens – fill Name: 'Ahmed Ali'",
            "Fill Phone: '+966501234567'",
            "Fill Message: 'I am interested in viewing this property'",
            "Click Send / Submit",
        ],
        "expected": "Modal opens with name, phone, and message fields. Submission shows success confirmation or closes modal.",
        "notes": "",
    },
    {
        "id": "TC-032", "feature": "Property – Save from detail",
        "title": "Save button in detail page saves property to list",
        "pre": "User is signed in",
        "steps": [
            "On property detail page, click the 'Save' button",
            "Navigate to /saved",
        ],
        "expected": "Property appears in the saved list. Save button state changes to indicate it is saved.",
        "notes": "",
    },
    {
        "id": "TC-033", "feature": "Property – Ask AI",
        "title": "Ask AI button navigates to advisor with property context",
        "pre": "Property detail page loaded",
        "steps": [
            "Click 'Ask AI' button",
        ],
        "expected": "Navigates to /advisor page. Optionally the property name/context is pre-loaded or referenced in the advisor.",
        "notes": "",
    },
    {
        "id": "TC-034", "feature": "Property – Comparable listings",
        "title": "Comparable listings section shows 3+ similar properties",
        "pre": "Multiple published properties in same or nearby district",
        "steps": [
            "Scroll to 'Comparable Listings' section",
            "Count the number of comparable property cards shown",
            "Click on one comparable card",
        ],
        "expected": "Section shows at least 2–3 comparable properties. Clicking opens their detail page.",
        "notes": "",
    },
]))

# ── S05: Compare ───────────────────────────────────────────────────────────────
sections.append(("Property Comparison", [
    {
        "id": "TC-035", "feature": "Compare – Page load",
        "title": "Compare page loads with pre-selected properties",
        "pre": "2 properties selected via search compare bar",
        "steps": [
            "From search, select 2 properties using Compare button",
            "Click 'Compare now'",
        ],
        "expected": "Navigate to /compare. Both properties shown as columns with images, scores, and price.",
        "notes": "",
    },
    {
        "id": "TC-036", "feature": "Compare – Add third property",
        "title": "User can add a third property to compare",
        "pre": "Compare page with 2 properties",
        "steps": [
            "Click the 'Add property' or '+' slot for the 3rd column",
            "Select a property from the dropdown/list",
        ],
        "expected": "Third property column added. Comparison tables update to show all 3 properties.",
        "notes": "Max 3 properties enforced",
    },
    {
        "id": "TC-037", "feature": "Compare – Remove property",
        "title": "Clicking X removes a property from comparison",
        "pre": "Compare page with 3 properties",
        "steps": [
            "Click the X / Remove button on one of the property columns",
        ],
        "expected": "That property column is removed. Remaining 2 properties stay. No page errors.",
        "notes": "",
    },
    {
        "id": "TC-038", "feature": "Compare – Financial table",
        "title": "Financial comparison table shows correct values",
        "pre": "Compare page with 2+ properties",
        "steps": [
            "Scroll to the Financial comparison section",
            "Check Annual Rent, Security Deposit, and Price per m² for each property",
            "Verify the best value is highlighted",
        ],
        "expected": "Annual rent matches property prices. Security deposit = rent × 2. Price/m² = rent ÷ area. Best (lowest) value highlighted in green.",
        "notes": "",
    },
    {
        "id": "TC-039", "feature": "Compare – Area scores table",
        "title": "Area comparison shows live scores from platform",
        "pre": "Properties in different districts. AreaIntelligence data seeded.",
        "steps": [
            "Scroll to Area comparison table",
            "Note Area Score, School Score, Family Score for each property",
        ],
        "expected": "Scores populated from live AreaIntelligence data. Best score highlighted per row.",
        "notes": "",
    },
    {
        "id": "TC-040", "feature": "Compare – AI recommendation",
        "title": "AI recommendation card picks the top property",
        "pre": "2+ properties on compare page",
        "steps": [
            "Scroll to AI Recommendation card at the top",
            "Read the winner selection and explanation",
        ],
        "expected": "AI card shows a winner with composite score, explanation paragraph, and highlights (cheapest, largest, best family fit). Winner property title is correct.",
        "notes": "",
    },
    {
        "id": "TC-041", "feature": "Compare – Save shortlist",
        "title": "Save shortlist button saves all compared properties",
        "pre": "User is signed in. 2+ properties on compare page.",
        "steps": [
            "Click 'Save shortlist' button",
            "Navigate to /saved",
        ],
        "expected": "All compared properties added to saved list. Confirmation shown before navigation.",
        "notes": "",
    },
]))

# ── S06: Saved Properties ──────────────────────────────────────────────────────
sections.append(("Saved Properties", [
    {
        "id": "TC-042", "feature": "Saved – Page load",
        "title": "Saved page shows user's saved properties",
        "pre": "User signed in with at least 2 saved properties",
        "steps": [
            "Navigate to http://localhost:8080/saved",
            "Count and verify property cards",
        ],
        "expected": "Saved page loads. Shows all previously saved properties with image, title, location, score, and saved date.",
        "notes": "",
    },
    {
        "id": "TC-043", "feature": "Saved – Filter by status",
        "title": "Status filter chips narrow displayed properties",
        "pre": "Saved page with multiple properties in different statuses",
        "steps": [
            "Click 'Inquiry sent' filter chip",
            "Verify only inquiry-sent properties shown",
            "Click 'All' chip to reset",
        ],
        "expected": "Filter chips work correctly. 'All' shows all properties regardless of status.",
        "notes": "",
    },
    {
        "id": "TC-044", "feature": "Saved – Add note",
        "title": "User can add a text note to a saved property",
        "pre": "Saved page loaded with at least one property",
        "steps": [
            "Click on a saved property card",
            "Find the 'Your Notes' section",
            "Type a note: 'Good location near school'",
            "Click Add / Submit",
        ],
        "expected": "Note appears in the notes list immediately. Note is persisted (still visible after page refresh).",
        "notes": "",
    },
    {
        "id": "TC-045", "feature": "Saved – Remove note",
        "title": "User can remove a note from a saved property",
        "pre": "Saved property with at least one note",
        "steps": [
            "Hover over an existing note",
            "Click the remove (X) button on the note",
        ],
        "expected": "Note disappears from the list immediately. Removal is persisted after page refresh.",
        "notes": "",
    },
    {
        "id": "TC-046", "feature": "Saved – Update inquiry status",
        "title": "Inquiry status stepper advances through steps",
        "pre": "Saved page with a property in 'Not contacted' status",
        "steps": [
            "On a saved property card, find the inquiry status tracker",
            "Click 'Mark as: Sent' (or equivalent step 1 button)",
            "Verify status updates to 'Inquiry sent'",
            "Click step 2: 'Contacted'",
            "Click step 3: 'Viewing Scheduled'",
            "When step 3 selected, set a viewing date",
        ],
        "expected": "Status updates at each step. 'Viewing Scheduled' step shows a datetime input. Status badge on card updates accordingly.",
        "notes": "",
    },
    {
        "id": "TC-047", "feature": "Saved – Remove property",
        "title": "Remove button deletes property from saved list",
        "pre": "Saved page with at least 2 properties",
        "steps": [
            "On a saved property card, click 'Remove' button",
            "Confirm removal if prompted",
        ],
        "expected": "Property card disappears from the list. Count in header decreases by 1. Removal is persisted after refresh.",
        "notes": "",
    },
    {
        "id": "TC-048", "feature": "Saved – Compare from saved",
        "title": "Compare all button navigates to compare with saved properties",
        "pre": "2+ saved properties",
        "steps": [
            "Click 'Compare all' button at the top of saved page",
        ],
        "expected": "Navigates to /compare with saved properties pre-loaded (up to 3).",
        "notes": "",
    },
]))

# ── S07: AI Advisor ─────────────────────────────────────────────────────────────
sections.append(("AI Advisor Chat", [
    {
        "id": "TC-049", "feature": "AI – Page load",
        "title": "AI Advisor page opens with empty chat and suggestions",
        "pre": "Portal running",
        "steps": [
            "Navigate to http://localhost:8080/advisor",
        ],
        "expected": "Page shows Maskan AI Advisor heading, 4 suggested question buttons, and an empty message area. No hardcoded conversation.",
        "notes": "",
    },
    {
        "id": "TC-050", "feature": "AI – Send message",
        "title": "User can type and send a question",
        "pre": "Advisor page loaded",
        "steps": [
            "Click in the text input box",
            "Type: 'What properties are available in Al Narjis?'",
            "Press Enter or click the send button",
            "Wait for response (5–15 seconds)",
        ],
        "expected": "User message appears right-aligned. Typing indicator (3 dots) shown while loading. AI response appears left-aligned with markdown formatting (headers, bold, bullet points).",
        "notes": "Response must reference REAL platform data, not generic LLM answer",
    },
    {
        "id": "TC-051", "feature": "AI – Platform data grounding",
        "title": "AI response contains actual property data from the platform",
        "pre": "Published properties and AreaIntelligence data in DB",
        "steps": [
            "Send: 'Which mediators are available on the platform?'",
            "Wait for response",
            "Check if response mentions actual license numbers or agency names from DB",
        ],
        "expected": "AI lists mediators with real license numbers, agency names, and phone numbers from the platform. Does NOT invent data.",
        "notes": "This is the core platform-grounding requirement",
    },
    {
        "id": "TC-052", "feature": "AI – Area comparison",
        "title": "AI compares districts using actual platform scores",
        "pre": "Area intelligence data for Al Narjis and Al Yasmin",
        "steps": [
            "Send: 'Compare Al Narjis vs Al Yasmin for a family'",
            "Read the AI response",
            "Cross-check scores mentioned against data in /areas page",
        ],
        "expected": "AI response references actual area scores (e.g. 'Al Narjis scores 65/100'). Values match what's shown in the Areas Intelligence console.",
        "notes": "",
    },
    {
        "id": "TC-053", "feature": "AI – Suggested questions",
        "title": "Clicking a suggested question sends it automatically",
        "pre": "Advisor page with empty chat",
        "steps": [
            "Click the first suggested question button",
            "Wait for response",
        ],
        "expected": "Question sent to AI. User message appears. AI responds within 15 seconds.",
        "notes": "",
    },
    {
        "id": "TC-054", "feature": "AI – Multi-turn conversation",
        "title": "AI maintains context across multiple messages",
        "pre": "Advisor page loaded",
        "steps": [
            "Send: 'Tell me about Al Narjis'",
            "After response, send: 'What about its rental prices?'",
            "After response, send: 'Is that good value?'",
        ],
        "expected": "AI understands that subsequent questions relate to Al Narjis from the conversation history. Does not ask for clarification unnecessarily.",
        "notes": "",
    },
    {
        "id": "TC-055", "feature": "AI – New chat button",
        "title": "New chat button clears the conversation",
        "pre": "Active conversation with multiple messages",
        "steps": [
            "Click 'New chat' button in the top bar",
        ],
        "expected": "All messages cleared. Empty state with suggestions shown again. Previous history does not appear.",
        "notes": "",
    },
    {
        "id": "TC-056", "feature": "AI – Textarea auto-expand",
        "title": "Input textarea grows when typing long text",
        "pre": "Advisor page loaded",
        "steps": [
            "Click in the textarea",
            "Type a long multi-line message (press Shift+Enter for new lines)",
        ],
        "expected": "Textarea height expands to show content, up to a max of ~4 lines, then scrolls internally.",
        "notes": "",
    },
]))

# ── S08: Mediator Flow ─────────────────────────────────────────────────────────
sections.append(("Mediator Registration & Dashboard", [
    {
        "id": "TC-057", "feature": "Mediator – Registration start",
        "title": "Mediator registration page accessible and shows 3-step flow",
        "pre": "User is signed in (non-mediator)",
        "steps": [
            "Navigate to http://localhost:8080/mediator",
            "If prompted (no mediator profile), click 'Become a mediator' or equivalent",
            "Navigate to /mediator/register",
        ],
        "expected": "Registration page shows 3 steps: Profile → Subscribe → Active. Step 1 is active.",
        "notes": "",
    },
    {
        "id": "TC-058", "feature": "Mediator – Step 1 profile",
        "title": "Mediator profile step validates and submits correctly",
        "pre": "On /mediator/register, Step 1 active",
        "steps": [
            "Leave License Number empty, click Continue — verify error",
            "Enter License Number: REGA-99999",
            "Enter Agency Name: Test Agency",
            "Enter Phone: +966501112233",
            "Enter Bio: 'Experienced property mediator in Riyadh'",
            "Click 'Continue to payment'",
        ],
        "expected": "Validation error shown when license is empty. On valid submission, Step 2 (payment) is shown.",
        "notes": "",
    },
    {
        "id": "TC-059", "feature": "Mediator – Step 2 payment",
        "title": "Mock payment step shows plan details and advances",
        "pre": "On Step 2 of mediator registration",
        "steps": [
            "Verify plan details: SAR 99/month, Standard Mediator Plan",
            "Click 'Pay SAR 99 and activate'",
        ],
        "expected": "Plan details shown correctly. After clicking pay, Step 3 (Done) is shown with success confirmation.",
        "notes": "This is a mock payment — no real charge",
    },
    {
        "id": "TC-060", "feature": "Mediator – Step 3 done",
        "title": "Registration completion redirects to mediator dashboard",
        "pre": "On Step 3 of mediator registration",
        "steps": [
            "Click 'Go to my dashboard' button",
        ],
        "expected": "Navigates to /mediator dashboard. Dashboard shows active subscription status.",
        "notes": "",
    },
    {
        "id": "TC-061", "feature": "Mediator – Dashboard stats",
        "title": "Dashboard shows pending, active, and total lead counts",
        "pre": "Active mediator profile. Leads assigned in DB.",
        "steps": [
            "Navigate to /mediator",
            "Check the 3 stat cards: Pending leads, Active leads, Total leads accepted",
        ],
        "expected": "Counts match the actual lead data from the DB. Cards show correct numbers.",
        "notes": "",
    },
    {
        "id": "TC-062", "feature": "Mediator – Add service area",
        "title": "Mediator can add a new service area",
        "pre": "Mediator dashboard loaded",
        "steps": [
            "In the Covered Areas section, enter District: 'Al Olaya'",
            "Select City: 'Riyadh'",
            "Click 'Add area'",
        ],
        "expected": "Area 'Al Olaya, Riyadh' added to the areas list immediately.",
        "notes": "",
    },
    {
        "id": "TC-063", "feature": "Mediator – Remove service area",
        "title": "Mediator can remove a service area",
        "pre": "Mediator dashboard with at least 2 covered areas",
        "steps": [
            "Click the remove button (X or trash) next to one area",
        ],
        "expected": "Area removed from the list. Count decreases.",
        "notes": "",
    },
    {
        "id": "TC-064", "feature": "Mediator – Lead detail",
        "title": "Mediator can view a lead and see contact details",
        "pre": "Pending lead exists in mediator's dashboard",
        "steps": [
            "Click 'View & Accept' on a pending lead",
        ],
        "expected": "Lead detail page opens showing full requirements, customer contact info, budget, and an accept/reject option.",
        "notes": "",
    },
]))

# ── S09: Admin Panel ───────────────────────────────────────────────────────────
sections.append(("Admin Panel – Listings", [
    {
        "id": "TC-065", "feature": "Admin – Access control",
        "title": "Admin panel only accessible to admin users",
        "pre": "Portal running",
        "steps": [
            "Sign in as a regular (non-admin) user",
            "Navigate to http://localhost:8080/admin",
        ],
        "expected": "Non-admin user is redirected away or shown an access denied message.",
        "notes": "Admin email is mnaushad.fms@gmail.com per config",
    },
    {
        "id": "TC-066", "feature": "Admin – Listings load",
        "title": "Admin listings tab shows all properties with status",
        "pre": "Signed in as admin",
        "steps": [
            "Open http://localhost:8080/admin",
            "Ensure Listings tab is active",
            "Check the 4 stat cards at the top",
            "Verify the table shows properties with ID, title, city, district, status",
        ],
        "expected": "Stat cards show total, pending, published, rejected counts. Table lists all properties paginated 20 per page.",
        "notes": "",
    },
    {
        "id": "TC-067", "feature": "Admin – Search listings",
        "title": "Admin search input filters the property table",
        "pre": "Admin listings page loaded",
        "steps": [
            "Type a property title or city in the search input",
            "Observe table update",
        ],
        "expected": "Table filters in real-time (or on Enter) to show only matching properties.",
        "notes": "",
    },
    {
        "id": "TC-068", "feature": "Admin – Status filter",
        "title": "Status dropdown filters table by property status",
        "pre": "Admin listings page loaded with mixed-status properties",
        "steps": [
            "Open Status filter dropdown",
            "Select 'Published'",
            "Verify all rows show Published status",
            "Select 'Pending Approval'",
        ],
        "expected": "Table shows only properties matching the selected status.",
        "notes": "",
    },
    {
        "id": "TC-069", "feature": "Admin – Create new listing",
        "title": "Admin can create a new property listing",
        "pre": "Admin listings page loaded",
        "steps": [
            "Click 'New listing' button",
            "In the drawer, fill: Title='Test Villa Riyadh', City='Riyadh', District='Al Narjis'",
            "Fill: Monthly Rent=7000, Bedrooms=4, Bathrooms=3, Area=350",
            "Set Status to 'Published'",
            "Click 'Save'",
        ],
        "expected": "Drawer closes. New listing appears in the table with status 'Published'. It is also visible on the search page.",
        "notes": "",
    },
    {
        "id": "TC-070", "feature": "Admin – Edit listing",
        "title": "Admin can edit an existing property",
        "pre": "Admin listings page with at least one property",
        "steps": [
            "Click the row menu (⋮) on a property",
            "Click 'Edit'",
            "Change the title",
            "Click Save",
        ],
        "expected": "Drawer opens pre-populated with property data. After save, table reflects the updated title.",
        "notes": "",
    },
    {
        "id": "TC-071", "feature": "Admin – Approve listing",
        "title": "Admin can approve a pending listing",
        "pre": "At least one 'Pending Approval' property exists",
        "steps": [
            "Click the row menu on a Pending Approval property",
            "Click 'Approve & publish'",
        ],
        "expected": "Property status changes to 'Published'. Stat card counts update accordingly.",
        "notes": "",
    },
    {
        "id": "TC-072", "feature": "Admin – Suspend/Reject listing",
        "title": "Admin can suspend or reject a listing",
        "pre": "Published property exists",
        "steps": [
            "Click row menu on a Published property",
            "Click 'Suspend'",
            "Verify status changes to 'Suspended'",
            "Verify property no longer appears in public search",
        ],
        "expected": "Status changes to Suspended in admin table. Suspended property hidden from /search results.",
        "notes": "",
    },
    {
        "id": "TC-073", "feature": "Admin – Delete listing",
        "title": "Admin can delete a property listing",
        "pre": "Test listing exists (created in TC-069)",
        "steps": [
            "Click row menu on the test listing",
            "Click 'Delete'",
            "Confirm deletion",
        ],
        "expected": "Listing removed from table. Total count decreases. Property no longer accessible via direct URL.",
        "notes": "Only delete test data",
    },
    {
        "id": "TC-074", "feature": "Admin – Bulk approve",
        "title": "Admin can bulk-approve multiple pending listings",
        "pre": "2+ pending listings exist",
        "steps": [
            "Filter by 'Pending Approval'",
            "Check checkboxes on 2 listings",
            "Click 'Approve' in the bulk action bar",
        ],
        "expected": "Both listings change to Published. Pending count decreases by 2.",
        "notes": "",
    },
    {
        "id": "TC-075", "feature": "Admin – Pagination",
        "title": "Listing table pagination navigates pages",
        "pre": "More than 20 listings exist in DB",
        "steps": [
            "Scroll to bottom of admin table",
            "Click 'Next page' button",
        ],
        "expected": "Table shows next 20 listings. Page indicator updates. Previous page button appears.",
        "notes": "",
    },
]))

sections.append(("Admin Panel – Mediators & Leads", [
    {
        "id": "TC-076", "feature": "Admin – Mediators tab",
        "title": "Mediators table shows all registered mediators",
        "pre": "Admin panel, at least one mediator registered",
        "steps": [
            "In admin sidebar, click 'Mediators'",
            "View the table",
        ],
        "expected": "Table shows mediators with: license number, agency name, subscription status, covered areas, total leads, verification status.",
        "notes": "",
    },
    {
        "id": "TC-077", "feature": "Admin – Verify mediator",
        "title": "Admin can verify a mediator",
        "pre": "Unverified mediator in the table",
        "steps": [
            "Click 'Verify' button on a pending mediator",
        ],
        "expected": "Mediator's row shows 'Verified' badge. On mediator's own dashboard, verified status is reflected.",
        "notes": "",
    },
    {
        "id": "TC-078", "feature": "Admin – Leads tab",
        "title": "Leads table shows all platform leads",
        "pre": "At least one lead submitted via lead form",
        "steps": [
            "In admin sidebar, click 'Leads'",
            "View the table",
        ],
        "expected": "Table shows leads with: ID, customer name & phone, area, max budget, status, created date.",
        "notes": "",
    },
    {
        "id": "TC-079", "feature": "Admin – Change lead status",
        "title": "Admin can update lead status from the table",
        "pre": "Lead with status 'open' exists",
        "steps": [
            "In the Leads table, find an 'open' lead",
            "Click the status dropdown on that row",
            "Select 'in_progress'",
        ],
        "expected": "Lead status updates to 'in_progress' immediately in the table.",
        "notes": "",
    },
]))

# ── S10: Areas Intelligence ────────────────────────────────────────────────────
sections.append(("Area Intelligence Console", [
    {
        "id": "TC-080", "feature": "Areas – Page load",
        "title": "Areas page loads with all district scores",
        "pre": "Area intelligence data seeded",
        "steps": [
            "Navigate to http://localhost:8080/areas",
            "Wait for table to load",
            "Count the number of districts shown",
        ],
        "expected": "Page loads. Stat cards show district count, avg area score, avg family score, avg rent. Table lists all districts.",
        "notes": "",
    },
    {
        "id": "TC-081", "feature": "Areas – City filter",
        "title": "City filter chip narrows table to selected city",
        "pre": "Areas page loaded",
        "steps": [
            "Click 'Riyadh' city filter chip",
            "Verify all rows in table show Riyadh city",
        ],
        "expected": "Only Riyadh districts shown. Count chip on 'Riyadh' button matches table row count.",
        "notes": "",
    },
    {
        "id": "TC-082", "feature": "Areas – Search filter",
        "title": "District search input filters table rows",
        "pre": "Areas page loaded",
        "steps": [
            "Type 'Al Narjis' in the search input",
        ],
        "expected": "Table shows only rows matching 'Al Narjis'. Partial match works (e.g. 'Narj' finds Al Narjis).",
        "notes": "",
    },
    {
        "id": "TC-083", "feature": "Areas – Lifestyle tag filter",
        "title": "Lifestyle tag buttons filter by district category",
        "pre": "Areas page loaded",
        "steps": [
            "Click 'Family Friendly' lifestyle tag button",
            "Check rows visible in table",
        ],
        "expected": "Only districts tagged as 'Family Friendly' are shown.",
        "notes": "",
    },
    {
        "id": "TC-084", "feature": "Areas – Sort columns",
        "title": "Clicking column headers sorts the table",
        "pre": "Areas page loaded with multiple districts",
        "steps": [
            "Click 'Area Score' column header",
            "Note order (highest to lowest or vice versa)",
            "Click again for reverse sort",
        ],
        "expected": "Districts sort by area score. Sort direction indicator (arrow) shows. Second click reverses order.",
        "notes": "Also test Family Score and Avg Rent sort",
    },
    {
        "id": "TC-085", "feature": "Areas – District detail panel",
        "title": "Clicking a district opens detail panel with all tabs",
        "pre": "Areas page loaded",
        "steps": [
            "Click 'Open' on the Al Narjis row",
            "Verify Overview tab is default",
            "Click Rental Trends tab",
            "Click Schools tab",
            "Click Hospitals tab",
            "Click Market Notes tab",
        ],
        "expected": "Detail panel opens. Each tab loads relevant data: overview with scores/description, trends with chart, schools list, hospitals list, market notes.",
        "notes": "",
    },
    {
        "id": "TC-086", "feature": "Areas – Add market note",
        "title": "User can add a market note to a district",
        "pre": "District detail panel open on Market Notes tab",
        "steps": [
            "Type a note: 'New metro station planned for 2026'",
            "Click 'Add' button",
        ],
        "expected": "Note added to the list immediately with AI insight badge.",
        "notes": "",
    },
    {
        "id": "TC-087", "feature": "Areas – Add new district",
        "title": "New district form creates a record and shows in table",
        "pre": "Areas page loaded. User has edit access.",
        "steps": [
            "Click 'New district' button",
            "Fill District Name: 'Al Muraba'",
            "Fill City: 'Riyadh'",
            "Set Avg Rent: 90000",
            "Set Area Score to 75",
            "Set Family Score to 80",
            "Click 'Add district'",
        ],
        "expected": "Modal closes. New district 'Al Muraba' appears in the table with entered scores.",
        "notes": "",
    },
]))

# ── S11: Navigation & Cross-cutting ────────────────────────────────────────────
sections.append(("Navigation, Layout & Cross-Cutting", [
    {
        "id": "TC-088", "feature": "Navigation – Breadcrumbs",
        "title": "Breadcrumbs and back links navigate correctly",
        "pre": "On property detail page reached via search",
        "steps": [
            "Check for a '← Back to search' or breadcrumb link",
            "Click it",
        ],
        "expected": "Returns to /search (with filters preserved if possible).",
        "notes": "",
    },
    {
        "id": "TC-089", "feature": "Navigation – Mobile responsive",
        "title": "Navigation and page layout work on mobile viewport",
        "pre": "Browser DevTools open",
        "steps": [
            "Press F12, toggle device toolbar, select iPhone 14 (390×844)",
            "Navigate to /, /search, /advisor, /property/1",
            "Check nav collapses to hamburger or mobile menu",
            "Check content not overflowing horizontally",
        ],
        "expected": "All pages render correctly at mobile size. No horizontal scroll on content. Navigation accessible.",
        "notes": "",
    },
    {
        "id": "TC-090", "feature": "Performance – Page load speed",
        "title": "Key pages load within acceptable time",
        "pre": "Network tab open in DevTools",
        "steps": [
            "Navigate to /search and measure time to show first properties",
            "Navigate to /property/1 and measure time to full render",
        ],
        "expected": "Search page shows properties within 3 seconds. Property detail within 5 seconds. No requests returning 4xx/5xx.",
        "notes": "",
    },
    {
        "id": "TC-091", "feature": "Error handling – Backend down",
        "title": "Friendly error shown when backend is unavailable",
        "pre": "Backend temporarily stopped (or wrong port)",
        "steps": [
            "Open /search when backend is offline",
        ],
        "expected": "Page shows a user-friendly error message (e.g. 'Unable to load live listings') instead of a crash or blank screen.",
        "notes": "",
    },
    {
        "id": "TC-092", "feature": "Data – SAR currency formatting",
        "title": "All monetary values display in SAR with correct formatting",
        "pre": "Properties loaded on any page",
        "steps": [
            "Check property prices on /search, /property/$id, and /compare",
        ],
        "expected": "All prices shown as 'SAR X,XXX' with comma separators. Annual and monthly rents are consistent (annual = monthly × 12).",
        "notes": "",
    },
    {
        "id": "TC-093", "feature": "Data – Property Score range",
        "title": "Property Scores on all pages fall within 55–97",
        "pre": "Properties loaded",
        "steps": [
            "Check Property Score on /search cards, /property/$id detail, and /compare table",
        ],
        "expected": "All scores are integers between 55 and 97 inclusive. No score is 0, 100, or the old hardcoded 72+formula result.",
        "notes": "",
    },
    {
        "id": "TC-094", "feature": "Lead – Submit rental inquiry",
        "title": "User can submit a rental lead from property detail",
        "pre": "Property detail page loaded",
        "steps": [
            "Find the lead/inquiry section or button in the right sidebar",
            "Fill in required details (name, phone, requirements)",
            "Submit the inquiry",
        ],
        "expected": "Lead recorded in DB. Confirmation shown to user. Lead appears in admin /leads table.",
        "notes": "",
    },
    {
        "id": "TC-095", "feature": "Session – Auth persistence",
        "title": "User session persists across page refreshes",
        "pre": "User is signed in",
        "steps": [
            "Sign in successfully",
            "Press F5 to refresh the page",
            "Check if user is still logged in",
        ],
        "expected": "User remains logged in after page refresh. Token is stored in localStorage or a persistent cookie.",
        "notes": "",
    },
]))

# ═══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
cover_page(doc)

# Summary table
p = doc.add_paragraph()
r = p.add_run("TEST SUMMARY")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = C_DARK

total_tcs = sum(len(v) for _, v in sections)
summary_rows = [("Total Test Cases", str(total_tcs)),
                ("Test Sections", str(len(sections))),
                ("Portal URL", "http://localhost:8080"),
                ("Backend API", "http://localhost:8000/api"),
                ("Admin Email", "mnaushad.fms@gmail.com"),
                ("Test Status", "☐ In Progress   ☐ Complete")]

t = doc.add_table(rows=len(summary_rows), cols=2)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (k, v) in enumerate(summary_rows):
    bg = "F1F5F9" if i % 2 == 0 else "FFFFFF"
    set_cell_bg(t.rows[i].cells[0], bg)
    set_cell_bg(t.rows[i].cells[1], bg)
    t.rows[i].cells[0].width = Cm(5)
    t.rows[i].cells[1].width = Cm(10)
    cell_para(t.rows[i].cells[0], k, bold=True, size=10)
    cell_para(t.rows[i].cells[1], v, size=10)

doc.add_page_break()

for sec_num, (sec_title, cases) in enumerate(sections, 1):
    section_banner(doc, sec_num, sec_title)
    make_table(doc, cases)
    if sec_num < len(sections):
        doc.add_page_break()

# ─── Save ─────────────────────────────────────────────────────────────────────
out_path = r"c:\Naushad\_AI Projects\_projects\Maskan Rental\Maskan_Portal_Test_Cases.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
print(f"  Sections  : {len(sections)}")
print(f"  Test cases: {total_tcs}")
