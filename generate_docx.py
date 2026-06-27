"""
generate_docx.py
Converts ARCHITECTURE.md → ARCHITECTURE.docx with professional formatting.
Run: python generate_docx.py
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ── Colours ──────────────────────────────────────────────────────────────────
C_BRAND    = RGBColor(0x0F, 0x43, 0x80)   # deep navy blue  – headings
C_ACCENT   = RGBColor(0x21, 0x6F, 0xDB)   # brand blue      – table headers
C_CODEBG   = RGBColor(0xF4, 0xF5, 0xF7)   # light grey      – code background
C_CODEFG   = RGBColor(0x1E, 0x1E, 0x1E)   # near-black      – code text
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_ROWALT   = RGBColor(0xF0, 0xF5, 0xFF)   # very light blue – alternate row
C_BORDER   = RGBColor(0xCC, 0xD9, 0xEE)   # table border

# ── Helpers ───────────────────────────────────────────────────────────────────
def rgb_hex(rgb: RGBColor) -> str:
    """Convert RGBColor (a tuple subclass) to 6-char uppercase hex string."""
    return "%02X%02X%02X" % (rgb[0], rgb[1], rgb[2])


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  rgb_hex(rgb))
    tcPr.append(shd)


def set_cell_border(cell, **edges):
    """edges: top/bottom/left/right, each a dict with keys: sz, color, val"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, conf in edges.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"),   conf.get("val",   "single"))
        el.set(qn("w:sz"),    str(conf.get("sz",   4)))
        el.set(qn("w:color"), conf.get("color", "auto"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def apply_all_cell_borders(cell, color="CCD9EE"):
    for edge in ("top", "left", "bottom", "right"):
        set_cell_border(cell, **{edge: {"val": "single", "sz": 4, "color": color}})


def para_shade(para, rgb: RGBColor):
    """Shade the paragraph background (for code blocks)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  rgb_hex(rgb))
    pPr.append(shd)


def add_run_formatted(para, text: str):
    """Parse inline **bold**, `code`, and plain text, adding runs."""
    # Split on **bold** and `code` patterns
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name  = "Courier New"
            run.font.size  = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)  # inline code red
        else:
            if part:
                para.add_run(part)


def set_para_spacing(para, before=0, after=0, line=None):
    pPr   = para._p.get_or_add_pPr()
    pSpacing = OxmlElement("w:spacing")
    pSpacing.set(qn("w:before"), str(before))
    pSpacing.set(qn("w:after"),  str(after))
    if line:
        pSpacing.set(qn("w:line"),     str(line))
        pSpacing.set(qn("w:lineRule"), "auto")
    pPr.append(pSpacing)


# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for s in doc.sections:
    s.top_margin    = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin   = Cm(2.5)
    s.right_margin  = Cm(2.5)

# Base Normal style
normal = doc.styles["Normal"]
normal.font.name  = "Calibri"
normal.font.size  = Pt(10.5)
normal.paragraph_format.space_after = Pt(4)

# Heading styles
for lvl, size, space_before in [(1, 18, 18), (2, 14, 14), (3, 11.5, 10)]:
    h = doc.styles[f"Heading {lvl}"]
    h.font.name        = "Calibri"
    h.font.size        = Pt(size)
    h.font.bold        = True
    h.font.color.rgb   = C_BRAND
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after  = Pt(4)
    h.paragraph_format.keep_with_next = True

# ── Cover Page ────────────────────────────────────────────────────────────────
def add_cover(doc):
    # Big title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=800, after=120)
    r = p.add_run("MASKAN RENTAL")
    r.font.name  = "Calibri"
    r.font.size  = Pt(32)
    r.font.bold  = True
    r.font.color.rgb = C_BRAND

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p2, before=0, after=240)
    r2 = p2.add_run("Architecture & Developer Reference")
    r2.font.name  = "Calibri"
    r2.font.size  = Pt(18)
    r2.font.color.rgb = C_ACCENT

    # Divider line
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("─" * 60)
    r3.font.color.rgb = C_BORDER
    r3.font.size = Pt(11)

    # Subtitle block
    for line, size, bold in [
        ("Saudi Rental Marketplace · Full-Stack Platform", 12, False),
        ("FastAPI · PostgreSQL · React · TanStack SSR · Anthropic AI", 11, False),
        ("Production-Ready · Horizontally Scalable · AWS Cloud Native", 10.5, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, before=60, after=60)
        r = p.add_run(line)
        r.font.name  = "Calibri"
        r.font.size  = Pt(size)
        r.font.bold  = bold
        r.font.color.rgb = RGBColor(0x44, 0x55, 0x77)

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p_date, before=240, after=0)
    r_d = p_date.add_run("Last Updated: June 2026")
    r_d.font.size  = Pt(10)
    r_d.font.color.rgb = RGBColor(0x77, 0x88, 0x99)

    doc.add_page_break()


add_cover(doc)

# ── Read source ───────────────────────────────────────────────────────────────
md_path = os.path.join(os.path.dirname(__file__), "ARCHITECTURE.md")
with open(md_path, encoding="utf-8") as f:
    lines = f.readlines()

# ── Tokenise into blocks ──────────────────────────────────────────────────────
# Block types: heading | code | table | hr | list | para
blocks = []
i = 0
while i < len(lines):
    line = lines[i].rstrip("\n")

    # Code fence
    if line.strip().startswith("```"):
        lang  = line.strip()[3:].strip()
        i    += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            code_lines.append(lines[i].rstrip("\n"))
            i += 1
        i += 1  # skip closing fence
        blocks.append({"type": "code", "lang": lang, "lines": code_lines})
        continue

    # Horizontal rule
    if re.match(r"^-{3,}$", line.strip()):
        blocks.append({"type": "hr"})
        i += 1
        continue

    # Heading
    m = re.match(r"^(#{1,4})\s+(.*)", line)
    if m:
        level = len(m.group(1))
        text  = m.group(2).strip()
        blocks.append({"type": "heading", "level": level, "text": text})
        i += 1
        continue

    # Table (line starts with |)
    if line.strip().startswith("|"):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_lines.append(lines[i].rstrip("\n"))
            i += 1
        blocks.append({"type": "table", "lines": table_lines})
        continue

    # List item
    m_li = re.match(r"^(\s*)[-*]\s+(.*)", line)
    m_nl = re.match(r"^(\s*)\d+\.\s+(.*)", line)
    if m_li or m_nl:
        list_items = []
        while i < len(lines):
            l = lines[i].rstrip("\n")
            m_li2 = re.match(r"^(\s*)[-*]\s+(.*)", l)
            m_nl2 = re.match(r"^(\s*)\d+\.\s+(.*)", l)
            if m_li2:
                indent = len(m_li2.group(1)) // 2
                list_items.append({"text": m_li2.group(2), "indent": indent, "ordered": False})
                i += 1
            elif m_nl2:
                indent = len(m_nl2.group(1)) // 2
                list_items.append({"text": m_nl2.group(2), "indent": indent, "ordered": True})
                i += 1
            else:
                break
        blocks.append({"type": "list", "items": list_items})
        continue

    # Blockquote
    if line.strip().startswith(">"):
        bq_lines = []
        while i < len(lines) and lines[i].strip().startswith(">"):
            bq_lines.append(lines[i].strip().lstrip(">").strip())
            i += 1
        blocks.append({"type": "blockquote", "lines": bq_lines})
        continue

    # Empty line
    if not line.strip():
        i += 1
        continue

    # Regular paragraph — collect until blank line or structural element
    para_lines = []
    while i < len(lines):
        l = lines[i].rstrip("\n")
        if not l.strip():
            break
        if l.strip().startswith(("```", "|", "#")) or re.match(r"^-{3,}$", l.strip()):
            break
        if re.match(r"^(\s*)[-*]\s+", l) or re.match(r"^(\s*)\d+\.\s+", l):
            break
        if l.strip().startswith(">"):
            break
        para_lines.append(l)
        i += 1
    if para_lines:
        blocks.append({"type": "para", "lines": para_lines})
    continue

# ── Render blocks → DOCX ─────────────────────────────────────────────────────
def parse_table_cells(row_line):
    """Split a markdown table row into cells, stripping leading/trailing |"""
    row_line = row_line.strip()
    if row_line.startswith("|"):
        row_line = row_line[1:]
    if row_line.endswith("|"):
        row_line = row_line[:-1]
    return [c.strip() for c in row_line.split("|")]


def is_separator_row(cells):
    return all(re.match(r"^:?-+:?$", c) for c in cells if c)


def render_code_block(doc, code_lines, lang=""):
    """Render a fenced code block as shaded monospace paragraphs."""
    # Top padding paragraph
    p_top = doc.add_paragraph()
    set_para_spacing(p_top, before=80, after=0)
    para_shade(p_top, C_CODEBG)

    if lang:
        p_lang = doc.add_paragraph()
        set_para_spacing(p_lang, before=0, after=0)
        para_shade(p_lang, RGBColor(0xD8, 0xE5, 0xF5))
        r = p_lang.add_run(f"  {lang.upper()}")
        r.font.name  = "Courier New"
        r.font.size  = Pt(7.5)
        r.font.bold  = True
        r.font.color.rgb = C_BRAND

    for cl in code_lines:
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=0)
        para_shade(p, C_CODEBG)
        # Preserve leading spaces by non-breaking space trick
        text = cl if cl.strip() else " "
        r = p.add_run(text)
        r.font.name  = "Courier New"
        r.font.size  = Pt(8.5)
        r.font.color.rgb = C_CODEFG

    # Bottom padding
    p_bot = doc.add_paragraph()
    set_para_spacing(p_bot, before=0, after=100)
    para_shade(p_bot, C_CODEBG)


def render_table(doc, table_lines):
    """Render a markdown table as a Word table."""
    if not table_lines:
        return

    all_rows = [parse_table_cells(l) for l in table_lines]
    # Filter out separator rows
    data_rows = [r for r in all_rows if not is_separator_row(r)]
    if not data_rows:
        return

    # Determine column count
    ncols = max(len(r) for r in data_rows)
    # Pad short rows
    data_rows = [r + [""] * (ncols - len(r)) for r in data_rows]

    tbl = doc.add_table(rows=len(data_rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # Remove default style borders — we'll set our own
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)

    for row_idx, row_data in enumerate(data_rows):
        row = tbl.rows[row_idx]
        is_header = row_idx == 0

        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]

            # Background
            if is_header:
                set_cell_bg(cell, C_ACCENT)
            elif row_idx % 2 == 0:
                set_cell_bg(cell, C_ROWALT)
            else:
                set_cell_bg(cell, C_WHITE)

            # Border
            apply_all_cell_borders(cell)

            # Padding
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement("w:tcMar")
            for side in ("top", "left", "bottom", "right"):
                m = OxmlElement(f"w:{side}")
                m.set(qn("w:w"),    "80")
                m.set(qn("w:type"), "dxa")
                tcMar.append(m)
            tcPr.append(tcMar)

            # Vertical alignment
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Text
            para = cell.paragraphs[0]
            para.clear()
            add_run_formatted(para, cell_text)

            # Formatting
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9.5)
                if is_header:
                    run.font.bold        = True
                    run.font.color.rgb   = C_WHITE

    # Set column widths (equal distribution)
    page_width_twips = 8640  # ~15.24 cm content width in twips
    col_width = int(page_width_twips / ncols)
    for col in tbl.columns:
        for cell in col.cells:
            cell.width = col_width * 20  # emu-ish; word adjusts

    doc.add_paragraph()  # spacing after table


def render_blockquote(doc, bq_lines):
    for bql in bq_lines:
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        # Left indent
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "400")
        pPr.append(ind)
        # Left border
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"),   "single")
        left.set(qn("w:sz"),    "18")
        left.set(qn("w:color"), "216FDB")
        pBdr.append(left)
        pPr.append(pBdr)
        set_para_spacing(p, before=40, after=40)

        r = p.add_run(bql)
        r.font.name      = "Calibri"
        r.font.size      = Pt(10)
        r.font.italic    = True
        r.font.color.rgb = RGBColor(0x44, 0x55, 0x77)


seen_h1 = False

for block in blocks:
    btype = block["type"]

    # ── Heading ──────────────────────────────────────────────────────────────
    if btype == "heading":
        lvl   = block["level"]
        text  = block["text"]
        # Strip markdown link syntax from headings: [text](#anchor)
        text  = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        # Remove trailing anchor tags
        text  = re.sub(r'\s*\{#[^}]+\}', '', text)

        if lvl == 1:
            if seen_h1:
                doc.add_page_break()
            seen_h1 = True
            p = doc.add_heading(text, level=1)
        elif lvl == 2:
            p = doc.add_heading(text, level=2)
        elif lvl == 3:
            p = doc.add_heading(text, level=3)
        else:
            p = doc.add_heading(text, level=3)

    # ── Horizontal rule ───────────────────────────────────────────────────────
    elif btype == "hr":
        p = doc.add_paragraph()
        set_para_spacing(p, before=60, after=60)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:color"), "CCD9EE")
        pBdr.append(bot)
        pPr.append(pBdr)

    # ── Code block ────────────────────────────────────────────────────────────
    elif btype == "code":
        render_code_block(doc, block["lines"], block.get("lang", ""))

    # ── Table ─────────────────────────────────────────────────────────────────
    elif btype == "table":
        render_table(doc, block["lines"])

    # ── Blockquote ────────────────────────────────────────────────────────────
    elif btype == "blockquote":
        render_blockquote(doc, block["lines"])

    # ── List ──────────────────────────────────────────────────────────────────
    elif btype == "list":
        for item in block["items"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Inches(0.25 + item["indent"] * 0.25)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.space_before = Pt(1)
            add_run_formatted(p, item["text"])
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"

    # ── Paragraph ─────────────────────────────────────────────────────────────
    elif btype == "para":
        full_text = " ".join(block["lines"]).strip()
        if not full_text:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.space_before = Pt(1)
        add_run_formatted(p, full_text)
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)


# ── Footer ────────────────────────────────────────────────────────────────────
from docx.oxml.ns import nsmap

section = doc.sections[0]
footer  = section.footer
fp = footer.paragraphs[0]
fp.clear()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("Maskan Rental — Architecture Reference  |  Confidential  |  2026")
r.font.size  = Pt(8)
r.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), "Maskan_Architecture_Reference_v2.docx")
doc.save(out_path)
print(f"Done. Saved: {out_path}")
